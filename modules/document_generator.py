import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

BASE_DIR = Path(__file__).parent.parent
DOCS_DIR = BASE_DIR / "generated_documents"
OFFER_DIR = DOCS_DIR / "offer_letters"
CERT_DIR = DOCS_DIR / "certificates"

# Ensure output directories exist
OFFER_DIR.mkdir(parents=True, exist_ok=True)
CERT_DIR.mkdir(parents=True, exist_ok=True)

def set_cell_border(cell, **kwargs):
    """Set cell borders for python-docx tables."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def generate_offer_letter(candidate_data):
    """
    Generate Offer Letter in both .docx and .pdf format.
    Returns (docx_filename, docx_path_str)
    """
    c_id = str(candidate_data.get("candidate_id", "C000")).strip()
    name = str(candidate_data.get("name", "Candidate")).strip()
    email = str(candidate_data.get("email", "")).strip()
    phone = str(candidate_data.get("phone", "")).strip()
    position = str(candidate_data.get("position", "")).strip()
    department = str(candidate_data.get("department", "")).strip()
    company = str(candidate_data.get("company", "ABC Technologies")).strip()
    joining_date = str(candidate_data.get("joining_date", "")).strip()
    salary = f"${float(candidate_data.get('salary', 0)):,.2f}"
    current_date = datetime.now().strftime("%B %d, %Y")

    # 1. Generate DOCX
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_comp = p_header.add_run(f"{company.upper()}\n")
    run_comp.bold = True
    run_comp.font.size = Pt(16)
    run_comp.font.color.rgb = RGBColor(30, 41, 59)

    run_sub = p_header.add_run("Human Resources Division | Corporate Headquarters\n")
    run_sub.font.size = Pt(9)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    run_date = p_header.add_run(f"Date: {current_date}\n")
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("OFFER OF EMPLOYMENT")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()

    p_to = doc.add_paragraph()
    p_to.add_run(f"To,\n").bold = True
    p_to.add_run(f"{name}\n").bold = True
    p_to.add_run(f"Candidate ID: {c_id}\n")
    p_to.add_run(f"Email: {email}\n")
    if phone:
        p_to.add_run(f"Phone: {phone}\n")

    doc.add_paragraph()

    p_body = doc.add_paragraph()
    p_body.paragraph_format.line_spacing = 1.2
    run_salut = p_body.add_run(f"Dear {name},\n\n")
    run_salut.bold = True

    p_body.add_run(
        f"Congratulations! On behalf of {company}, we are pleased to offer you the position of "
        f"{position} in our {department} department. We believe your experience and skills will be "
        f"a valuable asset to our organization.\n\n"
        f"Below are the primary details of your offer:"
    )

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    details = [
        ("Position Title:", position),
        ("Department:", department),
        ("Company Name:", company),
        ("Proposed Joining Date:", joining_date),
        ("Annual Salary / Compensation:", salary)
    ]

    for idx, (label, val) in enumerate(details):
        row = table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]

        cell_lbl.text = label
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)
        
        cell_val.text = str(val)
        cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(37, 99, 235)
        cell_val.paragraphs[0].runs[0].font.bold = True

        set_cell_border(cell_lbl, bottom={"val": "single", "sz": "4", "color": "CBD5E1"})
        set_cell_border(cell_val, bottom={"val": "single", "sz": "4", "color": "CBD5E1"})

    doc.add_paragraph()

    p_close = doc.add_paragraph()
    p_close.paragraph_format.line_spacing = 1.2
    p_close.add_run(
        f"Please review this document and indicate your acceptance.\n"
        f"We look forward to welcoming you to the team.\n\n"
        f"Sincerely,\n\n"
    )

    p_sign = doc.add_paragraph()
    p_sign.add_run("Authorized Recruitment Team\n").bold = True
    p_sign.add_run(f"Human Resources Department\n{company}")

    docx_filename = f"{c_id}_Offer_Letter.docx"
    docx_path = OFFER_DIR / docx_filename
    doc.save(docx_path)

    # 2. Generate PDF using ReportLab
    pdf_filename = f"{c_id}_Offer_Letter.pdf"
    pdf_path = OFFER_DIR / pdf_filename
    try:
        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=18, textColor=colors.HexColor('#2563EB'), alignment=1, spaceAfter=15)
        header_style = ParagraphStyle('HeaderStyle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, textColor=colors.HexColor('#1E293B'), alignment=2)
        sub_header_style = ParagraphStyle('SubHeaderStyle', parent=styles['Normal'], fontName='Helvetica', fontSize=9, textColor=colors.HexColor('#64748B'), alignment=2)
        body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#0F172A'))
        bold_body = ParagraphStyle('BoldBody', parent=body_style, fontName='Helvetica-Bold')

        elements = []
        elements.append(Paragraph(company.upper(), header_style))
        elements.append(Paragraph("Human Resources Division", sub_header_style))
        elements.append(Paragraph(f"Date: {current_date}", sub_header_style))
        elements.append(Spacer(1, 15))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CBD5E1'), spaceAfter=20))
        elements.append(Paragraph("OFFER OF EMPLOYMENT", title_style))
        elements.append(Spacer(1, 10))

        to_text = f"<b>To:</b><br/><b>{name}</b><br/>Candidate ID: {c_id}<br/>Email: {email}"
        if phone:
            to_text += f"<br/>Phone: {phone}"
        elements.append(Paragraph(to_text, body_style))
        elements.append(Spacer(1, 15))

        body_text = f"Dear <b>{name}</b>,<br/><br/>Congratulations! On behalf of {company}, we are pleased to offer you the position of <b>{position}</b> in our {department} department. We believe your experience will be a valuable asset to our organization.<br/><br/>Below are your offer details:"
        elements.append(Paragraph(body_text, body_style))
        elements.append(Spacer(1, 15))

        tbl_data = [
            [Paragraph("<b>Position Title:</b>", body_style), Paragraph(f"<b>{position}</b>", body_style)],
            [Paragraph("<b>Department:</b>", body_style), Paragraph(f"<b>{department}</b>", body_style)],
            [Paragraph("<b>Company Name:</b>", body_style), Paragraph(f"<b>{company}</b>", body_style)],
            [Paragraph("<b>Proposed Joining Date:</b>", body_style), Paragraph(f"<b>{joining_date}</b>", body_style)],
            [Paragraph("<b>Annual Salary:</b>", body_style), Paragraph(f"<b>{salary}</b>", body_style)]
        ]
        t = Table(tbl_data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 20))

        closing_text = f"Please indicate your acceptance of this offer.<br/>We look forward to working with you.<br/><br/>Sincerely,<br/><br/><b>Authorized HR Recruitment Team</b><br/>{company}"
        elements.append(Paragraph(closing_text, body_style))

        pdf_doc.build(elements)
    except Exception as e:
        print(f"PDF creation warning: {e}")

    return docx_filename, str(docx_path)

def generate_certificate(candidate_data):
    """
    Generate Certificate in both .docx and .pdf format.
    Returns (docx_filename, docx_path_str)
    """
    c_id = str(candidate_data.get("candidate_id", "C000")).strip()
    name = str(candidate_data.get("name", "Candidate")).strip()
    position = str(candidate_data.get("position", "")).strip()
    department = str(candidate_data.get("department", "")).strip()
    company = str(candidate_data.get("company", "ABC Technologies")).strip()
    joining_date = str(candidate_data.get("joining_date", "")).strip()
    current_date = datetime.now().strftime("%B %d, %Y")

    # 1. Generate DOCX
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cname = p_comp.add_run(f"{company.upper()}\n")
    run_cname.bold = True
    run_cname.font.size = Pt(20)
    run_cname.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph()

    p_cert = doc.add_paragraph()
    p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cert = p_cert.add_run("CERTIFICATE OF SELECTION")
    run_cert.bold = True
    run_cert.font.size = Pt(24)
    run_cert.font.color.rgb = RGBColor(37, 99, 235)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("THIS IS TO CERTIFY THAT")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(26)
    run_name.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph()

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.line_spacing = 1.3
    
    run_d1 = p_desc.add_run("has successfully completed the selection process and has been chosen for the position of\n")
    run_d1.font.size = Pt(13)
    
    run_pos = p_desc.add_run(f"{position}\n")
    run_pos.bold = True
    run_pos.font.size = Pt(16)
    run_pos.font.color.rgb = RGBColor(37, 99, 235)

    run_d2 = p_desc.add_run(f"in the {department} Department at {company}.\n\n")
    run_d2.font.size = Pt(13)

    run_jd = p_desc.add_run(f"Candidate ID: {c_id}   |   Scheduled Joining Date: {joining_date}\n")
    run_jd.font.size = Pt(11)
    run_jd.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    p_l = cell_left.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.add_run(f"Date of Issuance: {current_date}\nVerification Code: CERT-{c_id}-2026").font.size = Pt(10)

    p_r = cell_right.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_s1 = p_r.add_run("_________________________\nAuthorized HR Signatory\n")
    run_s1.bold = True
    run_s1.font.size = Pt(11)
    run_s2 = p_r.add_run(f"{company}")
    run_s2.font.size = Pt(10)

    docx_filename = f"{c_id}_Certificate.docx"
    docx_path = CERT_DIR / docx_filename
    doc.save(docx_path)

    # 2. Generate PDF using ReportLab
    pdf_filename = f"{c_id}_Certificate.pdf"
    pdf_path = CERT_DIR / pdf_filename
    try:
        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()

        c_header = ParagraphStyle('CHeader', fontName='Helvetica-Bold', fontSize=22, textColor=colors.HexColor('#1E293B'), alignment=1)
        c_title = ParagraphStyle('CTitle', fontName='Helvetica-Bold', fontSize=24, textColor=colors.HexColor('#2563EB'), alignment=1, spaceBefore=15)
        c_sub = ParagraphStyle('CSub', fontName='Helvetica', fontSize=12, textColor=colors.HexColor('#64748B'), alignment=1, spaceBefore=10)
        c_name = ParagraphStyle('CName', fontName='Helvetica-Bold', fontSize=26, textColor=colors.HexColor('#0F172A'), alignment=1, spaceBefore=15, spaceAfter=15)
        c_desc = ParagraphStyle('CDesc', fontName='Helvetica', fontSize=12, leading=18, textColor=colors.HexColor('#334155'), alignment=1)

        elements = []
        elements.append(Paragraph(company.upper(), c_header))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("CERTIFICATE OF SELECTION", c_title))
        elements.append(Paragraph("THIS IS TO CERTIFY THAT", c_sub))
        elements.append(Spacer(1, 10))
        elements.append(Paragraph(name, c_name))
        elements.append(Spacer(1, 10))
        
        desc_text = f"has successfully completed the selection process and has been chosen for the position of<br/><font size=14 color='#2563EB'><b>{position}</b></font><br/>in the {department} Department at {company}.<br/><br/><font size=10 color='#64748B'>Candidate ID: {c_id} | Scheduled Joining Date: {joining_date}</font>"
        elements.append(Paragraph(desc_text, c_desc))
        elements.append(Spacer(1, 40))

        tbl_data = [
            [Paragraph(f"Date of Issuance: {current_date}<br/>Code: CERT-{c_id}-2026", styles['Normal']),
             Paragraph(f"_______________________<br/><b>Authorized HR Signatory</b><br/>{company}", ParagraphStyle('RightSign', parent=styles['Normal'], alignment=2))]
        ]
        t = Table(tbl_data, colWidths=[250, 250])
        elements.append(t)

        pdf_doc.build(elements)
    except Exception as e:
        print(f"PDF certificate creation warning: {e}")

    return docx_filename, str(docx_path)
